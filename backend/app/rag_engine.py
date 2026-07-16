"""
RAG Engine for the Wildcat Navigator.

Supports two modes:
  - Dev mode  : ChromaDB + sentence-transformers (all-MiniLM-L6-v2), mock LLM
  - Prod mode : ChromaDB + AWS Bedrock embeddings + Claude 3.5 Sonnet
"""
from __future__ import annotations

import logging
import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Language detection
# ---------------------------------------------------------------------------

# Common Spanish words/phrases that reliably indicate Spanish input
_SPANISH_INDICATORS = [
    r"\bcómo\b", r"\bdónde\b", r"\bqué\b", r"\bcuál\b", r"\bcuándo\b",
    r"\bquién\b", r"\bpor\s+favor\b", r"\bgracias\b", r"\bnecesito\b",
    r"\bquiero\b", r"\bpuedo\b", r"\bhay\b", r"\bestoy\b", r"\btengo\b",
    r"\bpuede\b", r"\bdónde\b", r"\bcómo\s+puedo\b", r"\bcómo\s+me\b",
    r"\bestudiant[eo]\b", r"\buniversidad\b", r"\bcampus\b.*\bservicios\b",
    r"\bservicio[s]?\b", r"\bapartamento\b", r"\bestacionamiento\b",
    r"\bcomida\b", r"\bcafetería\b", r"\baccesibilidad\b", r"\bayuda\b",
    r"\bclases\b", r"\bhorario\b", r"\bmatrícula\b", r"\bbeca\b",
    r"\bcomo\b", r"\bdonde\b", r"\bque\b", r"\bpara\b.*\bestudiant",
]

def detect_language(text: str) -> str:
    """
    Detect whether the input is Spanish or English.
    Returns 'es' for Spanish, 'en' for English.
    Uses lightweight keyword matching — no external library needed.
    """
    text_lower = text.lower()
    for pattern in _SPANISH_INDICATORS:
        if re.search(pattern, text_lower):
            return "es"
    return "en"


# Workflow intent patterns
# ---------------------------------------------------------------------------
_WORKFLOW_PATTERNS: Dict[str, List[str]] = {
    "facility_rental": [
        r"\brent\b", r"\brental\b", r"\bbook\s+a\s+(room|venue|space|facility)\b",
        r"\breserve\s+a\s+(room|venue|space|facility)\b", r"\bfacility\b",
        r"\bvenue\b", r"\bevent\s+space\b",
    ],
    "accommodations": [
        r"\baccommodation", r"\bdisability\b", r"\barc\b",
        r"\baccessibility\b", r"\bextended\s+time\b", r"\btest\s+taking\b",
        r"\bnote.?taking\b", r"\b504\b", r"\biep\b",
    ],
    "parking_permit": [
        r"\bpark(ing)?\s+permit\b", r"\bbuy\s+a\s+permit\b",
        r"\bpurchase\s+a\s+permit\b", r"\bparking\s+pass\b",
        r"\bhow\s+do\s+i\s+get\s+a\s+park", r"\bpermit\b.*\bpark",
    ],
    "event_registration": [
        r"\bregister\s+(an?\s+)?event\b", r"\bevent\s+registration\b",
        r"\bhost\s+an?\s+event\b", r"\bplan\s+an?\s+event\b",
        r"\borganize\s+an?\s+event\b", r"\bcampus\s+event\b",
    ],
}

# ---------------------------------------------------------------------------
# Mock knowledge base (used when no documents are loaded)
# ---------------------------------------------------------------------------
_MOCK_DOCUMENTS = [
    {
        "id": "mock-001",
        "text": (
            "CSU Chico Admissions: To apply to California State University, Chico, "
            "students must submit an application through Cal State Apply at "
            "calstate.edu/apply. The application window for fall semester opens in "
            "October. Required materials include transcripts, test scores (optional "
            "for many programs), and personal insight questions."
        ),
        "metadata": {
            "title": "Admissions Overview",
            "url": "https://www.csuchico.edu/admissions/",
            "type": "website",
            "source": "mock",
        },
    },
    {
        "id": "mock-002",
        "text": (
            "Financial Aid at CSU Chico: Students can apply for financial aid by "
            "completing the FAFSA (fafsa.ed.gov) or California Dream Act Application. "
            "The priority deadline is March 2nd each year. The Financial Aid office "
            "is located in Student Services Center 250 and can be reached at "
            "(530) 898-6451 or finaid@csuchico.edu."
        ),
        "metadata": {
            "title": "Financial Aid",
            "url": "https://www.csuchico.edu/fa/",
            "type": "website",
            "source": "mock",
        },
    },
    {
        "id": "mock-003",
        "text": (
            "Housing & Residence Life: On-campus housing at CSU Chico includes "
            "Whitney Hall, Shasta Hall, Mechoopda Hall, and the Konkow Residence "
            "Complex. Applications open in January for the following academic year. "
            "Contact Housing at (530) 898-6204 or housing@csuchico.edu. "
            "The Housing office is in the BMU Room 101."
        ),
        "metadata": {
            "title": "Campus Housing",
            "url": "https://www.csuchico.edu/housing/",
            "type": "website",
            "source": "mock",
        },
    },
    {
        "id": "mock-004",
        "text": (
            "Registrar's Office: The Registrar manages course registration, "
            "transcripts, graduation applications, and enrollment verification. "
            "Students register for classes through the Student Center portal. "
            "Priority registration dates are assigned based on units completed. "
            "Office location: Kendall Hall 220. Phone: (530) 898-5142. "
            "Email: registrar@csuchico.edu."
        ),
        "metadata": {
            "title": "Office of the Registrar",
            "url": "https://www.csuchico.edu/registrar/",
            "type": "website",
            "source": "mock",
        },
    },
    {
        "id": "mock-005",
        "text": (
            "Student Health Center: The Student Health Services (SHS) provides "
            "primary care, mental health counseling, nutrition services, and "
            "health education programs to enrolled students. Located at Student "
            "Services Center 190. Phone: (530) 898-6452. After-hours nurse line "
            "available. Appointments can be scheduled online through the patient portal."
        ),
        "metadata": {
            "title": "Student Health Services",
            "url": "https://www.csuchico.edu/shs/",
            "type": "website",
            "source": "mock",
        },
    },
    {
        "id": "mock-006",
        "text": (
            "Parking Services at CSU Chico: Students may purchase semester parking "
            "permits online at parking.csuchico.edu. Permit types include general "
            "student lots (Lots 1, 2, 3, 6, 7), motorcycle, and evening/weekend permits. "
            "Daily permits are available at dispensers in each lot. "
            "Contact Parking at (530) 898-5475 or parking@csuchico.edu."
        ),
        "metadata": {
            "title": "University Parking Services",
            "url": "https://www.csuchico.edu/parking/",
            "type": "website",
            "source": "mock",
        },
    },
    {
        "id": "mock-007",
        "text": (
            "Accessibility Resource Center (ARC): ARC provides academic accommodations "
            "for students with disabilities including physical, learning, psychiatric, "
            "and chronic health conditions. Students must self-identify and provide "
            "documentation from a licensed provider. Located at Student Services Center "
            "170. Phone: (530) 898-5959. Email: arc@csuchico.edu."
        ),
        "metadata": {
            "title": "Accessibility Resource Center",
            "url": "https://www.csuchico.edu/arc/",
            "type": "website",
            "source": "mock",
        },
    },
    {
        "id": "mock-008",
        "text": (
            "Library Services: Meriam Library offers research assistance, study rooms, "
            "interlibrary loans, and digital resource access. Hours vary by semester. "
            "Research consultations can be booked with a subject librarian online. "
            "Phone: (530) 898-6501. The library is located centrally on campus "
            "near the BMU."
        ),
        "metadata": {
            "title": "Meriam Library",
            "url": "https://www.csuchico.edu/library/",
            "type": "website",
            "source": "mock",
        },
    },
]


# ---------------------------------------------------------------------------
# RAGEngine
# ---------------------------------------------------------------------------

class RAGEngine:
    """
    Retrieval-Augmented Generation engine.

    In dev mode (DEV_MODE=true or Bedrock not configured):
      - Uses ChromaDB with a local persistent store
      - Embeds text with sentence-transformers all-MiniLM-L6-v2
      - Generates answers by summarising retrieved chunks (no external LLM call)

    In prod mode:
      - Uses ChromaDB with Bedrock Titan embeddings
      - Sends retrieved context to Claude 3.5 Sonnet via AWS Bedrock
    """

    def __init__(self) -> None:
        from app.config import get_settings
        self.settings = get_settings()
        self._collection = None          # chromadb.Collection
        self._embed_fn = None            # callable(texts) -> list[list[float]]
        self._bedrock_client = None      # boto3 client or None
        self._initialised = False

    # ------------------------------------------------------------------
    # Startup
    # ------------------------------------------------------------------

    def initialize(self) -> None:
        """Set up ChromaDB collection and embedding function."""
        if self._initialised:
            return

        self._setup_chroma()

        if self.settings.dev_mode or not self.settings.bedrock_configured:
            self._setup_dev_embeddings()
            logger.info("RAGEngine started in DEV mode (sentence-transformers).")
        else:
            self._setup_bedrock()
            logger.info("RAGEngine started in PROD mode (AWS Bedrock).")

        self._load_documents()
        self._initialised = True

    def _setup_chroma(self) -> None:
        # Patch sqlite3 with pysqlite3-binary if the system sqlite is too old
        # (ChromaDB requires sqlite >= 3.35.0; some servers have older versions)
        try:
            import sqlite3
            if sqlite3.sqlite_version_info < (3, 35, 0):
                import sys
                __import__("pysqlite3")
                sys.modules["sqlite3"] = sys.modules.pop("pysqlite3")
        except ImportError:
            pass  # pysqlite3-binary not installed; proceed and let chromadb raise a clear error

        import chromadb

        persist_dir = Path(self.settings.chroma_persist_dir)
        persist_dir.mkdir(parents=True, exist_ok=True)

        client = chromadb.PersistentClient(path=str(persist_dir))
        self._collection = client.get_or_create_collection(
            name=self.settings.chroma_collection_name,
            metadata={"hnsw:space": "cosine"},
        )
        logger.info("ChromaDB collection '%s' ready at %s",
                    self.settings.chroma_collection_name, persist_dir)

    def _setup_dev_embeddings(self) -> None:
        from sentence_transformers import SentenceTransformer

        model = SentenceTransformer(self.settings.embedding_model_name)

        def _embed(texts: List[str]) -> List[List[float]]:
            return model.encode(texts, show_progress_bar=False).tolist()

        self._embed_fn = _embed
        logger.info("Embedding model '%s' loaded.", self.settings.embedding_model_name)

    def _setup_bedrock(self) -> None:
        import boto3
        import json

        session = boto3.Session(
            aws_access_key_id=self.settings.aws_access_key_id,
            aws_secret_access_key=self.settings.aws_secret_access_key,
            region_name=self.settings.aws_region,
        )
        self._bedrock_client = session.client("bedrock-runtime")

        def _embed(texts: List[str]) -> List[List[float]]:
            embeddings = []
            for text in texts:
                body = json.dumps({"inputText": text})
                response = self._bedrock_client.invoke_model(
                    modelId=self.settings.bedrock_embedding_model_id,
                    body=body,
                    contentType="application/json",
                    accept="application/json",
                )
                result = json.loads(response["body"].read())
                embeddings.append(result["embedding"])
            return embeddings

        self._embed_fn = _embed
        logger.info("Bedrock embedding model '%s' configured.",
                    self.settings.bedrock_embedding_model_id)

    # ------------------------------------------------------------------
    # Document loading
    # ------------------------------------------------------------------

    def _load_documents(self) -> None:
        """Index documents from the knowledge base directory (or mock data)."""
        kb_path = Path(self.settings.knowledge_base_dir)

        if not kb_path.exists() or not any(kb_path.iterdir()):
            logger.info("Knowledge base directory empty — loading mock documents.")
            self._index_documents(_MOCK_DOCUMENTS)
            return

        docs = []
        for file_path in kb_path.rglob("*"):
            if file_path.suffix.lower() in {".txt", ".md"}:
                docs.extend(self._parse_text_file(file_path))
            elif file_path.suffix.lower() == ".pdf":
                docs.extend(self._parse_pdf_file(file_path))
            elif file_path.suffix.lower() in {".docx", ".doc"}:
                docs.extend(self._parse_docx_file(file_path))

        if docs:
            self._index_documents(docs)
            logger.info("Indexed %d chunks from %s.", len(docs), kb_path)
        else:
            logger.info("No parseable files found — using mock documents.")
            self._index_documents(_MOCK_DOCUMENTS)

    def _parse_text_file(self, path: Path) -> List[dict]:
        """Split a text/markdown file into ~500-character chunks."""
        text = path.read_text(encoding="utf-8", errors="ignore")
        chunks = self._chunk_text(text, chunk_size=500, overlap=50)
        doc_type = "policy" if "policy" in path.stem.lower() else "website"
        return [
            {
                "id": f"{path.stem}-{i}",
                "text": chunk,
                "metadata": {
                    "title": path.stem.replace("_", " ").title(),
                    "url": "",
                    "type": doc_type,
                    "source": str(path),
                },
            }
            for i, chunk in enumerate(chunks)
        ]

    def _parse_pdf_file(self, path: Path) -> List[dict]:
        """Parse a PDF file into chunks (requires pypdf; skip gracefully if absent)."""
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.warning("pypdf not installed; skipping %s", path)
            return []

        reader = PdfReader(str(path))
        text = " ".join(
            page.extract_text() or "" for page in reader.pages
        )
        chunks = self._chunk_text(text, chunk_size=500, overlap=50)
        return [
            {
                "id": f"{path.stem}-pdf-{i}",
                "text": chunk,
                "metadata": {
                    "title": path.stem.replace("_", " ").title(),
                    "url": "",
                    "type": "pdf",
                    "source": str(path),
                },
            }
            for i, chunk in enumerate(chunks)
        ]

    def _parse_docx_file(self, path: Path) -> List[dict]:
        """Parse a Word (.docx) file into chunks (requires python-docx; skip gracefully if absent)."""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed; skipping %s. Run: pip install python-docx", path)
            return []

        try:
            doc = Document(str(path))
            # Extract all paragraph text, skip empty lines
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also pull text from tables — keep each row as a complete unit
            table_rows = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_rows.append(row_text)
        except Exception as exc:
            logger.warning("Could not parse Word document %s: %s", path, exc)
            return []

        # For documents that are primarily tables (like building links),
        # chunk by grouping complete rows together instead of splitting mid-row
        if table_rows and len(table_rows) > len(paragraphs):
            # Table-heavy document: chunk by rows, never splitting a row
            chunks = self._chunk_by_lines(table_rows, max_chunk_size=1200)
        else:
            # Paragraph-heavy document: use standard text chunking
            text = "\n".join(paragraphs)
            if table_rows:
                text += "\n" + "\n".join(table_rows)
            chunks = self._chunk_text(text, chunk_size=500, overlap=50)
        return [
            {
                "id": f"{path.stem}-docx-{i}",
                "text": chunk,
                "metadata": {
                    "title": path.stem.replace("_", " ").replace("-", " ").title(),
                    "url": "",
                    "type": "policy",
                    "source": str(path),
                },
            }
            for i, chunk in enumerate(chunks)
        ]

    @staticmethod
    def _chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping character-level chunks."""
        text = text.strip()
        if len(text) <= chunk_size:
            return [text] if text else []
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start += chunk_size - overlap
        return chunks

    @staticmethod
    def _chunk_by_lines(lines: List[str], max_chunk_size: int = 1200) -> List[str]:
        """Group complete lines into chunks without splitting any line.
        This prevents URLs and table rows from being cut mid-way."""
        if not lines:
            return []
        chunks = []
        current_chunk = []
        current_size = 0
        for line in lines:
            line_size = len(line) + 1  # +1 for newline
            if current_size + line_size > max_chunk_size and current_chunk:
                chunks.append("\n".join(current_chunk))
                current_chunk = []
                current_size = 0
            current_chunk.append(line)
            current_size += line_size
        if current_chunk:
            chunks.append("\n".join(current_chunk))
        return chunks

    def _index_documents(self, docs: List[dict]) -> None:
        """Upsert documents into the ChromaDB collection."""
        if not docs:
            return

        existing_ids = set(self._collection.get(include=[])["ids"])
        new_docs = [d for d in docs if d["id"] not in existing_ids]

        if not new_docs:
            logger.info("All %d documents already indexed.", len(docs))
            return

        texts = [d["text"] for d in new_docs]
        embeddings = self._embed_fn(texts)

        self._collection.upsert(
            ids=[d["id"] for d in new_docs],
            embeddings=embeddings,
            documents=texts,
            metadatas=[d["metadata"] for d in new_docs],
        )
        logger.info("Upserted %d new document chunks.", len(new_docs))

    # ------------------------------------------------------------------
    # Querying
    # ------------------------------------------------------------------

    def query(
        self,
        question: str,
        chat_history: Optional[List[dict]] = None,
    ) -> Tuple[str, List[dict], float, str]:
        """
        Retrieve relevant chunks and generate an answer.

        Args:
            question:     The user's question string.
            chat_history: Optional list of prior {role, content} dicts.

        Returns:
            (answer_text, source_docs, confidence_score, detected_language)
            where source_docs is a list of metadata dicts,
            confidence_score is a float in [0, 1], and
            detected_language is 'en' or 'es'.
        """
        if not self._initialised:
            self.initialize()

        # Detect language before querying
        language = detect_language(question)

        # For Spanish queries, also search with English translation of key terms
        # (embedding model is English-focused)
        search_text = question
        if language == "es":
            search_text = self._translate_es_keywords(question)

        # Embed the question
        q_embedding = self._embed_fn([search_text])[0]

        # Retrieve top-k chunks from ChromaDB
        results = self._collection.query(
            query_embeddings=[q_embedding],
            n_results=min(self.settings.top_k_results, self._collection.count() or 1),
            include=["documents", "metadatas", "distances"],
        )

        documents: List[str] = results["documents"][0] if results["documents"] else []
        metadatas: List[dict] = results["metadatas"][0] if results["metadatas"] else []
        distances: List[float] = results["distances"][0] if results["distances"] else []

        # Convert cosine distances to similarity scores [0, 1]
        similarities = [max(0.0, 1.0 - d) for d in distances]
        confidence = float(sum(similarities) / len(similarities)) if similarities else 0.0

        # Generate answer
        if self._bedrock_client is not None:
            answer = self._bedrock_answer(question, documents, chat_history or [], language)
        else:
            answer = self._build_mock_answer(question, documents, metadatas, language)

        return answer, metadatas, confidence, language

    # Common Spanish-to-English translations for search queries
    _ES_TO_EN = {
        "autobuses": "bus transit",
        "autobús": "bus",
        "transporte": "transportation",
        "horario": "schedule",
        "mapa": "map campus",
        "estacionamiento": "parking",
        "comida": "dining food",
        "comedor": "dining",
        "biblioteca": "library",
        "clases": "classes schedule",
        "carreras": "majors programs",
        "programas": "programs majors",
        "admisión": "admissions",
        "admisiones": "admissions",
        "costos": "cost tuition",
        "ayuda financiera": "financial aid",
        "becas": "scholarships",
        "vivienda": "housing",
        "dormitorios": "housing dorms",
        "seguridad": "safety security",
        "salud": "health wellness",
        "eventos": "events",
        "deportes": "athletics sports",
        "centro": "downtown",
        "información": "information",
        "dónde": "where",
        "cómo": "how",
        "cuándo": "when",
        "qué": "what",
        "aplicar": "apply admissions",
        "inscribir": "enroll register",
        "matrícula": "tuition enrollment",
        "campus": "campus",
        "universidad": "university",
        "estudiante": "student",
        "accesibilidad": "accessibility",
        "bicicleta": "bike bicycle",
        "caminar": "walk walking",
        "servicios": "services",
    }

    def _translate_es_keywords(self, question: str) -> str:
        """Translate Spanish keywords to English for better embedding search."""
        text = question.lower()
        translated_terms = []
        for es_term, en_term in self._ES_TO_EN.items():
            if es_term in text:
                translated_terms.append(en_term)
        if translated_terms:
            return " ".join(translated_terms)
        # Fallback: return original question (model may still find something)
        return question

    def _build_mock_answer(
        self,
        question: str,
        documents: List[str],
        metadatas: List[dict],
        language: str = "en",
    ) -> str:
        """
        Construct a contextual answer from retrieved chunks without an LLM.
        Responds in Spanish when language='es'.
        Cleans up partial sentences, formats URLs as markdown links, and
        presents information in a readable format.
        """
        import re

        if not documents:
            if language == "es":
                return (
                    "Lo siento, no tengo información específica sobre ese tema en mi "
                    "base de conocimientos en este momento. Por favor, comuníquese "
                    "directamente con el departamento correspondiente de CSU Chico para obtener ayuda."
                )
            return (
                "I'm sorry, I don't have specific information about that topic in my "
                "knowledge base right now. Please contact the relevant CSU Chico "
                "department directly for assistance."
            )

        # Special handler for map-focused queries — return a clean, focused response
        q_lower = question.lower()
        map_query_keywords = ['campus map', 'show me a map', 'map of campus', 'mapa del campus', 'mapa']
        if any(kw in q_lower for kw in map_query_keywords):
            building_links = self._get_building_links_section(language)
            if language == "es":
                return (
                    "🗺️ **Mapa del Campus CSU Chico**\n\n"
                    "![Mapa del Campus](/campus-map.png)\n\n"
                    "---\n\n"
                    "📄 **Mapa para imprimir (PDF):** [Descargar mapa del campus](https://www.csuchico.edu/_assets/documents/office/admissions/printable-campus-map.pdf)\n\n"
                    "♿ **Mapas de rutas accesibles:** [csuchico.edu/facilities/accessibility-maps](https://csuchico.edu/facilities/accessibility-maps)\n\n"
                    "🌐 **Mapa interactivo en línea:** [csuchico.edu/maps/campus](https://www.csuchico.edu/maps/campus/)"
                    + building_links
                )
            return (
                "🗺️ **CSU Chico Campus Map**\n\n"
                "![CSU Chico Campus Map](/campus-map.png)\n\n"
                "---\n\n"
                "📄 **Printable Map (PDF):** [Download printable campus map](https://www.csuchico.edu/_assets/documents/office/admissions/printable-campus-map.pdf)\n\n"
                "♿ **Accessibility Route Maps:** [csuchico.edu/facilities/accessibility-maps](https://csuchico.edu/facilities/accessibility-maps)\n\n"
                "🌐 **Interactive Online Map:** [csuchico.edu/maps/campus](https://www.csuchico.edu/maps/campus/)"
                + building_links
            )

        def _format_urls(text: str) -> str:
            """Convert raw URLs into readable markdown links."""
            # Pattern: "Label | apple_maps_url | google_maps_url" (from Buildings Links doc)
            def _format_building_line(line: str) -> str:
                # Split carefully — URLs don't contain bare pipes, so split on ' | '
                parts = [p.strip() for p in line.split(' | ')]
                if len(parts) < 2:
                    # Try splitting on single pipe with spaces
                    parts = [p.strip() for p in line.split('|')]
                
                # Separate URLs from names
                urls_in_parts = [p for p in parts if p.startswith('http')]
                names_in_parts = [p for p in parts if p and not p.startswith('http')]
                
                if len(urls_in_parts) >= 1 and len(names_in_parts) >= 1:
                    name = names_in_parts[0]
                    # Skip header rows
                    if name.lower() in ('building', 'name', 'location'):
                        return None
                    links = []
                    for url in urls_in_parts:
                        url = url.rstrip(',').strip()
                        if 'maps.apple.com' in url:
                            links.append(f"[Apple Maps]({url})")
                        elif 'google.com/maps' in url:
                            links.append(f"[Google Maps]({url})")
                        else:
                            links.append(f"[Link]({url})")
                    return f"- **{name}** — {' | '.join(links)}"
                return None

            lines = text.split('\n')
            formatted_lines = []
            for line in lines:
                # Skip table header rows
                if line.strip().lower().startswith('building') and 'apple maps' in line.lower() and 'google maps' in line.lower():
                    continue
                if '|' in line and 'http' in line:
                    formatted = _format_building_line(line)
                    if formatted:
                        formatted_lines.append(formatted)
                        continue

                # Convert standalone raw URLs to markdown links
                # Match URLs not already in markdown link syntax
                def _url_to_link(match):
                    url = match.group(0).rstrip(',')
                    # Try to create a readable label from the URL
                    if 'csuchico.edu' in url:
                        # Extract path for label
                        path = url.split('csuchico.edu')[-1].strip('/')
                        label = path.split('/')[-1].replace('-', ' ').replace('_', ' ').title() if path else 'CSU Chico'
                        if label.endswith('.pdf'):
                            label = label[:-4] + ' (PDF)'
                        return f"[{label}]({url})"
                    elif 'maps.apple.com' in url:
                        return f"[Apple Maps]({url})"
                    elif 'google.com/maps' in url:
                        return f"[Google Maps]({url})"
                    else:
                        return f"[Link]({url})"

                # Only replace URLs that aren't already part of a markdown link
                line = re.sub(
                    r'(?<!\()(https?://[^\s,\)]+)',
                    _url_to_link,
                    line
                )
                formatted_lines.append(line)

            return '\n'.join(formatted_lines)

        def _clean_snippet(text: str) -> str:
            """Trim to complete sentences and clean up fragments."""
            # Remove leading partial sentence (starts with lowercase or no capital)
            lines = text.strip().split('\n')
            # If first line looks like a fragment (doesn't start with uppercase, #, -, *, or **)
            if lines and lines[0] and not lines[0][0].isupper() and not lines[0][0] in '#-*|':
                # Find the first sentence boundary or line break
                first_period = lines[0].find('. ')
                first_newline = text.find('\n')
                cut = min(
                    first_period + 2 if first_period >= 0 else len(text),
                    first_newline if first_newline >= 0 else len(text)
                )
                text = text[cut:].strip()

            # Trim to last complete sentence (ends with . ? ! or a markdown line)
            if len(text) > 1500:
                text = text[:1800]
                # Find last sentence end or list item end
                for end_char in ['\n- ', '\n* ', '. ', '.\n', '?\n', '!\n', '? ', '! ']:
                    last = text.rfind(end_char)
                    if last > 800:
                        text = text[:last + 1]
                        break
                else:
                    # Fall back to last newline
                    last_nl = text.rfind('\n')
                    if last_nl > 800:
                        text = text[:last_nl]

            # Format URLs in the cleaned text
            text = _format_urls(text)
            return text.strip()

        # Group all retrieved chunks by source title so we can combine
        # fragments from the same document (e.g., a long building list)
        from collections import OrderedDict
        grouped: OrderedDict[str, List[str]] = OrderedDict()
        for doc, meta in zip(documents, metadatas):
            title = meta.get("title", "Campus Resource")
            if title not in grouped:
                grouped[title] = []
            grouped[title].append(doc)

        # Build body parts, combining chunks from the same source
        body_parts = []
        for title, chunks in grouped.items():
            # Combine all chunks from same source, removing overlap duplicates
            combined = chunks[0]
            for chunk in chunks[1:]:
                # Try to find overlap and merge
                overlap_len = min(60, len(combined), len(chunk))
                overlap_found = False
                for ol in range(overlap_len, 10, -1):
                    if combined.endswith(chunk[:ol]):
                        combined += chunk[ol:]
                        overlap_found = True
                        break
                if not overlap_found:
                    combined += "\n" + chunk

            snippet = _clean_snippet(combined)
            if not snippet:
                continue
            body_parts.append(f"**{title}**\n\n{snippet}")

            # Limit to 3 distinct sources to keep response manageable
            if len(body_parts) >= 3:
                break

        if not body_parts:
            # Fallback if cleaning removed everything
            snippet = _format_urls(documents[0][:300])
            title = metadatas[0].get("title", "Campus Resource")
            body_parts.append(f"**{title}**\n\n{snippet}")

        if language == "es":
            intro = "Esto es lo que encontré:\n\n"
        else:
            intro = "Here's what I found:\n\n"

        # Add source links if available
        sources_note = ""
        urls = [m.get("url", "") for m in metadatas if m.get("url")]
        if urls:
            if language == "es":
                sources_note = "\n\n---\n📎 Para más información: " + " | ".join(urls[:2])
            else:
                sources_note = "\n\n---\n📎 For more details: " + " | ".join(urls[:2])

        # Include campus map image if the question is about maps/directions
        map_keywords = ['map', 'mapa', 'where is', 'dónde', 'directions', 'find', 'locate', 'building']
        map_image = ""
        if any(kw in question.lower() for kw in map_keywords):
            if language == "es":
                map_image = "\n\n---\n\n🗺️ **Mapa del Campus:**\n\n![Mapa del Campus CSU Chico](/campus-map.png)"
            else:
                map_image = "\n\n---\n\n🗺️ **Campus Map:**\n\n![CSU Chico Campus Map](/campus-map.png)"

        # Always include building map links when user asks about buildings
        building_keywords = ['building', 'buildings', 'edificio', 'edificios', 'map link',
                            'where is', 'dónde', 'locate', 'find', 'directions',
                            'library', 'dining', 'union', 'recreation', 'admin',
                            'engineering', 'performing arts', 'kitchen']
        building_links_section = ""
        if any(kw in question.lower() for kw in building_keywords):
            # Check if building links are already in the response
            if 'Apple Maps' not in "\n".join(body_parts):
                building_links_section = self._get_building_links_section(language)

        return intro + "\n\n---\n\n".join(body_parts) + building_links_section + sources_note + map_image

    def _get_building_links_section(self, language: str = "en") -> str:
        """Return formatted building map links section."""
        from pathlib import Path
        try:
            from docx import Document
            kb_path = Path(self.settings.knowledge_base_dir)
            # Look for the buildings links file
            for docx_file in kb_path.glob("*[Bb]uilding*[Ll]ink*.docx"):
                doc = Document(str(docx_file))
                lines = []
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if len(cells) >= 3 and cells[0].lower() != 'building':
                            name = cells[0]
                            apple_url = cells[1]
                            google_url = cells[2]
                            lines.append(
                                f"- **{name}** — [Apple Maps]({apple_url}) | [Google Maps]({google_url})"
                            )
                if lines:
                    if language == "es":
                        header = "\n\n---\n\n📍 **Enlaces de Mapas de Edificios:**\n\n"
                    else:
                        header = "\n\n---\n\n📍 **Building Map Links:**\n\n"
                    return header + "\n".join(lines)
        except Exception:
            pass
        return ""

    def _bedrock_answer(
        self,
        question: str,
        documents: List[str],
        chat_history: List[dict],
        language: str = "en",
    ) -> str:
        """Send context + question to Claude via AWS Bedrock and return the answer."""
        import json

        context = "\n\n---\n\n".join(documents[:5])
        history_text = ""
        if chat_history:
            turns = []
            for msg in chat_history[-6:]:
                role = msg.get("role", "user").capitalize()
                turns.append(f"{role}: {msg.get('content', '')}")
            history_text = "\n".join(turns) + "\n\n"

        if language == "es":
            system_prompt = (
                "Eres el Conserje Virtual Wildcat, un asistente útil para la Universidad "
                "Estatal de California, Chico (CSU Chico). El estudiante escribe en español, "
                "por lo tanto DEBES responder completamente en español. "
                "Responde usando ÚNICAMENTE el contexto proporcionado. "
                "Sé conciso, amable y preciso. Si el contexto no contiene suficiente "
                "información, dilo y dirige al estudiante a la oficina universitaria apropiada."
            )
        else:
            system_prompt = (
                "You are the Wildcat Navigator, a helpful assistant for California State "
                "University, Chico (CSU Chico). Answer the student's question using ONLY the "
                "provided context. Be concise, friendly, and accurate. If the context does "
                "not contain enough information, say so and direct the student to the "
                "appropriate campus office."
            )

        user_message = (
            f"Context from CSU Chico knowledge base:\n{context}\n\n"
            f"{history_text}"
            f"Student question: {question}"
        )

        body = json.dumps({
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 1024,
            "system": system_prompt,
            "messages": [{"role": "user", "content": user_message}],
        })

        response = self._bedrock_client.invoke_model(
            modelId=self.settings.bedrock_model_id,
            body=body,
            contentType="application/json",
            accept="application/json",
        )
        result = json.loads(response["body"].read())
        return result["content"][0]["text"]

    # ------------------------------------------------------------------
    # Workflow intent detection
    # ------------------------------------------------------------------

    def detect_workflow_intent(self, question: str) -> Optional[str]:
        """
        Detect whether the question maps to a known workflow type.

        Returns the workflow key (e.g. 'facility_rental') or None.
        """
        q_lower = question.lower()
        for workflow_type, patterns in _WORKFLOW_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, q_lower):
                    logger.debug("Workflow intent '%s' detected.", workflow_type)
                    return workflow_type
        return None


# ---------------------------------------------------------------------------
# Module-level singleton
# ---------------------------------------------------------------------------

_engine_instance: Optional[RAGEngine] = None


def get_rag_engine() -> RAGEngine:
    """Return the module-level RAGEngine singleton (not yet initialised)."""
    global _engine_instance
    if _engine_instance is None:
        _engine_instance = RAGEngine()
    return _engine_instance
